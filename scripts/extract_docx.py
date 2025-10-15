#!/usr/bin/env python3
"""
Extract text content from DOCX files for analysis.
Used to understand contract structure for synthetic data generation.
"""

import sys
from pathlib import Path
from docx import Document
import json


def extract_docx_structure(docx_path: Path) -> dict:
    """
    Extract structured content from DOCX file.

    Returns:
        dict with sections, paragraphs, tables, and metadata
    """
    doc = Document(docx_path)

    result = {
        "filename": docx_path.name,
        "paragraphs": [],
        "tables": [],
        "styles_used": set(),
        "stats": {
            "total_paragraphs": 0,
            "total_tables": 0,
            "total_chars": 0,
        }
    }

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            result["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })
            result["styles_used"].add(para.style.name if para.style else "Normal")
            result["stats"]["total_chars"] += len(para.text)

    result["stats"]["total_paragraphs"] = len(result["paragraphs"])

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append(table_data)

    result["stats"]["total_tables"] = len(result["tables"])
    result["styles_used"] = sorted(list(result["styles_used"]))

    return result


def print_document_summary(data: dict):
    """Print readable summary of document structure."""
    print(f"\n{'='*80}")
    print(f"FILE: {data['filename']}")
    print(f"{'='*80}")
    print(f"\nSTATISTICS:")
    print(f"  Paragraphs: {data['stats']['total_paragraphs']}")
    print(f"  Tables: {data['stats']['total_tables']}")
    print(f"  Characters: {data['stats']['total_chars']:,}")
    print(f"  Styles used: {', '.join(data['styles_used'])}")

    print(f"\n{'─'*80}")
    print("DOCUMENT CONTENT:")
    print(f"{'─'*80}\n")

    for i, para in enumerate(data["paragraphs"], 1):
        # Show style if not Normal
        style_marker = f" [{para['style']}]" if para['style'] != "Normal" else ""
        print(f"{para['text']}{style_marker}\n")

    if data["tables"]:
        print(f"\n{'─'*80}")
        print("TABLES:")
        print(f"{'─'*80}\n")
        for i, table in enumerate(data["tables"], 1):
            print(f"Table {i}:")
            for row in table:
                print("  | " + " | ".join(row) + " |")
            print()


def identify_clauses(data: dict) -> list:
    """
    Identify potential clause sections based on content patterns.

    Returns:
        List of identified clauses with metadata
    """
    clauses = []

    for para in data["paragraphs"]:
        text = para["text"].strip()

        # Look for numbered sections or headings
        # Common patterns: "1. Title", "Section 1:", "Article 1", etc.
        if any([
            text.startswith(tuple(f"{i}." for i in range(1, 100))),
            "section" in text.lower()[:30],
            "article" in text.lower()[:30],
            para["style"] in ["Heading 1", "Heading 2", "Heading 3"],
        ]):
            clauses.append({
                "title": text[:100],  # First 100 chars
                "style": para["style"],
                "full_text": text,
            })

    return clauses


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)

    # Extract structure
    data = extract_docx_structure(docx_path)

    # Print summary
    print_document_summary(data)

    # Identify clauses
    clauses = identify_clauses(data)
    if clauses:
        print(f"\n{'─'*80}")
        print(f"IDENTIFIED CLAUSES: {len(clauses)}")
        print(f"{'─'*80}\n")
        for i, clause in enumerate(clauses, 1):
            print(f"{i}. {clause['title']}")
